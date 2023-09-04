# -*- coding: utf-8 -*-
"""
Created on Sat Sep  2 11:47:15 2023

@author: aryan
"""

# First, we import all the necessary library functions
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Function to apply formatting to a single document
def format_document(doc, font_name, font_size, is_bold, is_italic, alignment):
    # Convert font size to Pt (points)
    new_font_size = Pt(font_size)

    # Map alignment text to WD_PARAGRAPH_ALIGNMENT enum
    alignment_mapping = {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }

    # Iterate through paragraphs and runs in the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Apply formatting to the runs (text) in the paragraph
            run.font.name = font_name
            run.font.size = new_font_size
            run.font.bold = is_bold
            run.font.italic = is_italic

        # Set paragraph alignment
        paragraph.alignment = alignment_mapping.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)


# Ask the user for the folder location
folder_path = input("Enter the folder location containing the Word documents: ")

# Check if the specified folder exists
if not os.path.exists(folder_path): 
    print("The specified folder does not exist.")
else:
    # Load the configuration document for formatting settings
    config_doc = Document('C:\\Users\\aryan\\Downloads\\config.docx')

    # Initialize formatting settings with default values
    font_name = 'Arial'
    font_size = 12
    is_bold = False
    is_italic = False
    alignment = 'left'

    # Extract formatting settings from the configuration Word document
    for paragraph in config_doc.paragraphs:
        for run in paragraph.runs:
            text = run.text.strip()

            if text.startswith("<<FONT_SIZE:"):
                font_size = int(text.split(":")[1].strip())
            elif text.startswith("<<FONT_NAME:"):
                font_name = text.split(":")[1].strip()
            elif text.startswith("<<BOLD:"):
                is_bold = text.split(":")[1].strip().lower() == "true"
            elif text.startswith("<<ITALIC:"):
                is_italic = text.split(":")[1].strip().lower() == "true"
            elif text.startswith("<<ALIGNMENT:"):
                alignment = text.split(":")[1].strip().lower()

    # Loop through all files in the folder
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            # Open the document
            doc = Document(os.path.join(folder_path, filename))

            # Apply the formatting settings from the configuration document
            format_document(doc, font_name, font_size, is_bold, is_italic, alignment)

            # Save the modified document with the same filename
           # Save the modified document with the same filename
            new_filename = os.path.splitext(filename)[0] + '_formatted.docx'
            doc.save(os.path.join(folder_path, new_filename))

    print("Formatting complete.")